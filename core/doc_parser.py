from docx import Document
import pdfplumber
from utils.text_helper import clean_text

def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    full = ""
    # 段落
    for p in doc.paragraphs:
        full += p.text
    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full += cell.text
    return clean_text(full)

def read_pdf(file_path: str) -> str:
    full = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            full += page_text
    return clean_text(full)

def parse_contract(file_path: str) -> str:
    if file_path.lower().endswith(".docx"):
        return read_docx(file_path)
    elif file_path.lower().endswith(".pdf"):
        return read_pdf(file_path)
    else:
        raise Exception("仅支持docx、pdf格式合同文件")